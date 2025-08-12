# e2w/e2w.py

import os
import requests
import pandas as pd
from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from typing import List, Dict, Any, Optional
import gc
from .styles.page_layout import PageLayout, Orientation, Size
from .styles.font_family import FontFamily, FontStyle
from .styles.table_style import TableFormat

class ExportToWord:
    '''Support converting the template with html tag to Word file'''

    def __init__(self, 
                 template_content: str= None, 
                 template_file: str=None, 
                 context:dict= None, 
                 output_path: str="output.docx",
                 page_layout: PageLayout = PageLayout(Orientation.LANDSCAPE, Size.A4),
                 font_family: FontFamily = FontFamily(),
                 table_style: TableFormat = TableFormat(),
                 heading_levels: int = 6,
                 image_max_size: tuple = (5.3, 3.5),  # Default max size for images in inches
                 error_font: FontFamily = FontFamily(name="Arial", size=8, style=FontStyle.ITALIC, color=RGBColor(255,0,0)),
                 max_workers: int = 4,  # Number of threads for multithreading
                 chunk_size: int = 1000  # Memory management: process data in chunks
                 ):
        '''Initialize the E2W class with context, template path, and output path.'''
        self.template_content = ""
        self.context = context or {}
        self.max_workers = max_workers
        self.chunk_size = chunk_size
        self._lock = threading.Lock()  # Thread safety for document operations
        
        if template_file:
            if not context:
                raise ValueError("If temp_file is provided, context must not be None.")
            if not os.path.exists(template_file):
                raise FileNotFoundError(f"Template file not found: {template_file}")
            # read template file and binding data to content
            _template = self._clean_template(template_file)
            self.template_content =  self._replace_variables(_template)
        elif template_content:
            self.template_content = template_content
        else:
            raise ValueError("Either template_file or template_content must be provided.")
        self.template_content = self._format_template_to_html(self.template_content)
        self.output_path = output_path
        self.page_layout = page_layout
        self.font_family = font_family
        self.error_font = error_font
        self.page_dimensions = page_layout.size
        self.table_style = table_style
        self.heading_levels = heading_levels
        self.image_max_size = (Inches(image_max_size[0]), Inches(image_max_size[1]))  
        self.align_paragraph = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        self.document = self._document_setup()

    def render(self):
        '''Render word document from template with multithreading support'''
        soup = BeautifulSoup(self.template_content, 'html.parser')
        
        # Process tags in parallel for better performance
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Submit all tag processing tasks
            future_to_tag = {
                executor.submit(self._process_tag, tag): tag 
                for tag in soup.contents if isinstance(tag, Tag)
            }
            
            # Process completed tasks
            for future in as_completed(future_to_tag):
                try:
                    future.result()
                except Exception as e:
                    # Log error and continue with other tags
                    print(f"Error processing tag: {e}")
                    continue
        
        # Save the document to the specified output path
        self.document.save(self.output_path)
        
        # Clean up memory
        gc.collect()

    def _process_tag(self, tag: Tag):
        """Process a single tag with thread safety."""
        try:
            handler = self._tag_handlers().get(tag.name)
            if handler:
                with self._lock:
                    handler(tag)
            else:
                with self._lock:
                    self._add_paragraph(tag)
        except Exception as e:
            print(f"Error processing tag {tag.name}: {e}")

    def _format_template_to_html(self, content: str) -> str:
        """Convert content to HTML tag, properly handle empty lines and HTML tags."""
        lines = content.splitlines()
        html_lines = []
        in_html_block = False
        current_html_block = []
        
        for line in lines:
            stripped = line.strip()
            
            # Handle empty lines - preserve them as <br/> tags
            if not stripped:
                if in_html_block:
                    # If we're in an HTML block, add a line break
                    current_html_block.append("<br/>")
                else:
                    # If we're in text mode, add a paragraph break
                    html_lines.append("<p></p>")
                continue
            
            # Handle comments
            if stripped.startswith("#"):
                continue
            
            # Check if this line contains HTML tags
            if self._contains_html_tags(stripped):
                # If we were building a text block, close it first
                if not in_html_block and current_html_block:
                    html_lines.append(f"<p>{''.join(current_html_block)}</p>")
                    current_html_block = []
                
                # Process HTML line
                if self._is_complete_html_tag(stripped):
                    # Complete HTML tag on single line
                    html_lines.append(stripped)
                else:
                    # Multi-line HTML tag
                    in_html_block = True
                    current_html_block.append(stripped)
            else:
                # Text line
                if in_html_block:
                    # Close HTML block and start text
                    html_lines.append(''.join(current_html_block))
                    current_html_block = []
                    in_html_block = False
                
                # Add text as paragraph
                current_html_block.append(stripped)
        
        # Handle any remaining content
        if current_html_block:
            if in_html_block:
                html_lines.append(''.join(current_html_block))
            else:
                html_lines.append(f"<p>{''.join(current_html_block)}</p>")
        
        return "\n".join(html_lines)
    
    def _contains_html_tags(self, line: str) -> bool:
        """Check if a line contains HTML tags."""
        return '<' in line and '>' in line
    
    def _is_complete_html_tag(self, line: str) -> bool:
        """Check if a line contains a complete HTML tag."""
        # Simple heuristic: check if line starts and ends with tags
        # or contains balanced tags
        open_tags = line.count('<')
        close_tags = line.count('>')
        return open_tags == close_tags and open_tags > 0
        
    def _tag_handlers(self):
        """Mapping of tag names to handler methods."""
        return {
            'header': lambda tag : self._handle_header_footer(self.document.sections[0], tag),
            'footer': lambda tag: self._handle_header_footer(self.document.sections[0], tag),
            'title': self._handle_title,
            **{f'h{i}': self._handle_heading for i in range(1, self.heading_levels+1)},  # h1 to h6
            # 'image': lambda tag: self._handle_image(tag.get('src')) if tag.get('src') else None,
            'image': self._handle_image, 
            'dataframe': self._handle_dataframe,
            'base64-image': self._handle_base64_image,
            'session_break': self._handle_add_section, 
            'page_break': self._handle_add_page_break, 
            'p': self._add_paragraph,
            "blank_table": self._handle_blank_table,
            'ul': self._handle_list,
            'ol': self._handle_list,
        }
    

    def _document_setup(self):
        """Document setup such as page size, orientation, font family, font style, etc."""
        document = Document()
        section = document.sections[0]
        _width, _height = self.page_dimensions
        if self.page_layout.orientation == Orientation.LANDSCAPE.value:
            section.page_width = Inches(_height)
            section.page_height = Inches(_width)
            self.page_dimensions = (_height, _width)  # Update dimensions for landscape
        else:
            section.page_width = Inches(_width)
            section.page_height = Inches(_height)
        self._set_margins(section, top=1.0, bottom=1.0, left=1.25, right=1.25)

        # Set up the default font family for the document.
        style = document.styles[self.font_family.style.value]
        style.font.name = self.font_family.name
        style.font.size = self.font_family.size
        style.font.color.rgb = self.font_family.color
        return document
    
    def _set_margins(self, section, top: float = 1.0, bottom: float = 1.0, left: float = 1.25, right: float = 1.25):
        """Set the margins for the document section."""
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

    def _set_error_font_style(self, run):
        '''Change the font color to red and italicize it.'''
        run.font.name = self.error_font.name
        run.font.color.rgb = self.error_font.color
        run.font.size = self.error_font.size
        if self.error_font.style == FontStyle.ITALIC:
            run.font.italic = True
        elif self.error_font.style == FontStyle.BOLD:
            run.font.bold = True    


    def _clean_template(self, template_file: str):
        """Remove all comments in the template content. 
        That is, lines starting with #."""
        with open(template_file, 'r', encoding='utf-8') as f:
            _content = f.read()
        cleaned_lines = []
        for line in _content.split('\n'):
            stripped = line.strip()
            if not stripped.startswith('#'):
                cleaned_lines.append(line)
        return '\n'.join(cleaned_lines)
    

    def _handle_header_footer(self, section, tag): 
        """Process header and footer in the template."""
        _rows = 1
        _columns = 1
        _height = 0.5
        section_part = section.footer
        if tag.name == "header":
            section_part = section.header
            _columns = 3

        _table_width = Inches(self.page_dimensions[0] / _columns)
        table = section_part.add_table(rows=_rows, cols=_columns, width=Inches(self.page_dimensions[0]))
        table.autofit = True
        table.alignment = WD_ALIGN_VERTICAL.CENTER
        
        for column in table.columns:
            column.width = _table_width 
        _first_cell = table.rows[0].cells[0]
        _first_cell.text = tag.text.strip()
        _last_cell = table.rows[0].cells[-1]
        _image = tag.find('image')
        if _image and 'src' in _image.attrs:        
            self._handle_image(_image, aligment='right', 
                               paragraph=_last_cell.paragraphs[0], height=_height) 
            
    def _handle_image(self, tag: Tag, #image_path: str, 
                      aligment: str = 'center',
                      paragraph=None, 
                      height: float = 0.0):
        """Insert an image into the document with improved error handling and memory management."""
        _image_path = tag.get('src', None)
        if not _image_path:
            self._add_error_paragraph("Image tag must have a 'src' attribute.")
            return
            
        _align = tag.get('align', aligment).lower()
        _aligntment = self.align_paragraph.get(_align, WD_ALIGN_PARAGRAPH.CENTER)
        paragraph = self.document.add_paragraph() if paragraph is None else paragraph
        paragraph.alignment = _aligntment 
        _image_path = _image_path.strip()
        
        try:
            if os.path.exists(_image_path):
                _width, _height = self._get_image_size(_image_path, height) if height != 0.0 else self._get_image_size(_image_path)
                _tag_width, _tag_height = tag.get('width', None), tag.get('height', None)
                
                if _tag_width or _tag_height:
                    if _tag_width:
                        _width = Inches(float(_tag_width))
                    if _tag_height:
                        _height = Inches(float(_tag_height))
                elif _width > self.image_max_size[0] or _height > self.image_max_size[1]:
                    _width = min(_width, self.image_max_size[0])
                    _height = min(_height, self.image_max_size[1])
                
                run = paragraph.add_run()
                run.add_picture(_image_path, _width, _height)
                
                # Clean up memory after adding image
                gc.collect()
            else:
                run = paragraph.add_run(f"[Missing image: {_image_path}]")
                self._set_error_font_style(run)
        except Exception as e:
            error_msg = f"[Error loading image {_image_path}: {e}]"
            run = paragraph.add_run(error_msg)
            self._set_error_font_style(run)

    def _get_image_size(self, image_path: str, target_height: float = 0.0) -> tuple:
        """Calculates the target size for the image while preserving the aspect ratio with memory optimization."""
        try:
            from PIL import Image
            with Image.open(image_path) as img:
                img_width, img_height = img.size
                aspect_ratio = img_width / img_height
                
                if target_height == 0.0:
                    target_width = self.page_dimensions[0] * 0.6
                    target_height = target_width / aspect_ratio
                else:
                    target_width = target_height * aspect_ratio
                
                # Ensure dimensions don't exceed page limits
                max_width = self.page_dimensions[0] * 0.9
                max_height = self.page_dimensions[1] * 0.8
                
                if target_width > max_width:
                    target_width = max_width
                    target_height = target_width / aspect_ratio
                
                if target_height > max_height:
                    target_height = max_height
                    target_width = target_height * aspect_ratio
                
                return (Inches(target_width), Inches(target_height))
        except Exception as e:
            print(f"Error calculating image size: {e}")
            # Return default size on error
            return (Inches(4), Inches(3))

    def _handle_title(self, tag: Tag):
        '''Handle the title tag in the template.'''
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(tag.get_text().strip().upper())
        run.bold = True
        run.font.size = Pt(16)  
        
    def _handle_heading(self, tag: Tag):
        """Handle the heading tags in the template."""
        level = int(tag.name[1:-1]) if tag.name[1:-1].isdigit() else 1  # Extract number from h1, h2, etc.
        heading = self.document.add_heading(level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.add_run(tag.get_text().strip())

    def _handle_blank_table(self, tag: Tag):
        '''Handle the blank table tag in the template.'''
        """Create a blank table with the specified number of rows and columns."""
        _rows = int(tag.get('rows', 1))
        _columns = int(tag.get('columns', 1))
        _columns_name = tag.get('columns_name', '').split(',')
        _table_columns = max(_columns, len(_columns_name))
        table = self.document.add_table(rows=_rows, cols=_columns, style=self.table_style.style.value)
        table.autofit = True
        header_cells = table.rows[0].cells
        if len(_columns_name) < _table_columns:
            for i in range(len(_columns_name), _table_columns):
                _columns_name.append(f"Column {i+1}")
        for i, column in enumerate(_columns_name):
            if i < _table_columns:
                header_cells[i].text = column.strip()
                header_cells[i].paragraphs[0].runs[0].bold = True
        for _ in range(1, _rows):
            row_cells = table.add_row().cells
            for i in range(_table_columns):
                if i < len(_columns_name):
                    row_cells[i].text = ""
                    row_cells[i].paragraphs[0].runs[0].bold = False

    def _handle_dataframe(self, tag: Tag):
        """Handle the dataframe tag in the template with improved memory management and index removal."""
        df = pd.DataFrame()
        if 'src' in tag.attrs:
            if os.path.exists(tag['src']):
                try:
                    # Read CSV in chunks for memory efficiency
                    chunk_list = []
                    for chunk in pd.read_csv(tag['src'], chunksize=self.chunk_size):
                        chunk_list.append(chunk)
                    if chunk_list:
                        df = pd.concat(chunk_list, ignore_index=True)
                        del chunk_list  # Free memory
                        gc.collect()
                except Exception as e:
                    self._add_error_paragraph(f"Error reading CSV file: {e}")
                    return
        elif 'api' in tag.attrs:
            api_url = tag['api']
            api_config = self.context.get('apis',{}).get(api_url, {})
            headers = self.context.get('api_headers',{}) 
            params = api_config.get('params', {})
            try:
                resp = requests.post(api_url, json=params, headers=headers)
                if resp.status_code == 200:
                    data = resp.json()
                    if data:
                        df = pd.DataFrame(data.get('data',{}))
            except Exception as e:
                self._add_error_paragraph(f"API Error: {e}")
                return
        
        if df.empty:
            self._add_error_paragraph(f"{tag.get_text()} No data available.")
            return
        
        # Remove index column if it's the default pandas index
        if df.index.name is None and df.index.dtype == 'int64':
            df = df.reset_index(drop=True)
        
        # Get table attributes for customization
        table_style = tag.get('style', self.table_style.style.value)
        max_rows = int(tag.get('max_rows', 0))  # 0 means all rows
        
        # Limit rows if specified
        if max_rows > 0 and len(df) > max_rows:
            df = df.head(max_rows)
        
        # Create table with proper styling
        table = self.document.add_table(rows=1, cols=len(df.columns), style=table_style)
        table.autofit = True
        
        # Process header
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = str(column)
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Process data in chunks for memory efficiency
        total_rows = len(df)
        if total_rows > self.chunk_size:
            # Process large datasets in chunks with multithreading
            self._process_dataframe_chunks(df, table, total_rows)
        else:
            # Process small datasets directly
            self._process_dataframe_rows(df, table)
        
        # Clean up memory
        del df
        gc.collect()
    
    def _process_dataframe_chunks(self, df: pd.DataFrame, table, total_rows: int):
        """Process large dataframes in chunks with multithreading."""
        chunks = [df.iloc[i:i+self.chunk_size] for i in range(0, total_rows, self.chunk_size)]
        
        with ThreadPoolExecutor(max_workers=min(self.max_workers, len(chunks))) as executor:
            # Submit chunk processing tasks
            future_to_chunk = {
                executor.submit(self._process_chunk, chunk, table): chunk 
                for chunk in chunks
            }
            
            # Process completed chunks
            for future in as_completed(future_to_chunk):
                try:
                    future.result()
                except Exception as e:
                    print(f"Error processing dataframe chunk: {e}")
                    continue
    
    def _process_chunk(self, chunk: pd.DataFrame, table):
        """Process a single chunk of dataframe data."""
        with self._lock:
            for _, row in chunk.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value) if pd.notna(value) else ''
                    row_cells[i].paragraphs[0].runs[0].bold = False
    
    def _process_dataframe_rows(self, df: pd.DataFrame, table):
        """Process dataframe rows directly for small datasets."""
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ''
                row_cells[i].paragraphs[0].runs[0].bold = False
    
    def _add_error_paragraph(self, message: str):
        """Add an error message paragraph with error styling."""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(message)
        self._set_error_font_style(run)

    def _handle_add_section(self, tag: Tag):
        """Handle session break in the template."""
        section = self.document.add_section()
        self._set_margins(section)
    
    def _handle_add_page_break(self, tag: Tag):
        """Handle page break in the template."""
        self.document.add_page_break()
        self._set_margins(self.document.sections[-1])

    def _add_paragraph(self, tag):
        '''Add a paragraph from either a Tag or a plain text line.'''
        para = self.document.add_paragraph()
        if isinstance(tag, Tag):
            if 'align' in tag.attrs:
                para.alignment = self.align_paragraph.get(tag.get('align', '').lower(), WD_ALIGN_PARAGRAPH.LEFT)
            self._handle_inline_formatting(tag, para)
        elif isinstance(tag, str):
            _text = tag.strip()
            if _text:
                para.add_run(_text)

    def _handle_inline_formatting(self, element: Tag, para):
        """Recursively process inline tags <b>, <i>, <u> and normal text into the same paragraph."""
        def recurse(node, bold=False, italic=False, underline=False):
            if isinstance(node, NavigableString):
                text = str(node)
                if text.strip():
                    run = para.add_run(text)
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline
            elif isinstance(node, Tag):
                next_bold = bold or (node.name == 'b')
                next_italic = italic or (node.name == 'i')
                next_underline = underline or (node.name == 'u')
                for child in node.children:
                    recurse(child, next_bold, next_italic, next_underline)
        recurse(element)

    def _replace_variables(self, text: str) -> str:
        """Replace placeholders in the content with values from the context."""
        try:
            for key, value in self.context.items():
                placeholder = f"<{key}/>"
                if placeholder in text:
                    # Handle different data types
                    if isinstance(value, (list, tuple)):
                        # Convert lists to comma-separated strings
                        text = text.replace(placeholder, ', '.join(map(str, value)))
                    elif isinstance(value, dict):
                        # Convert dicts to formatted strings
                        formatted_dict = '\n'.join([f"{k}: {v}" for k, v in value.items()])
                        text = text.replace(placeholder, formatted_dict)
                    else:
                        text = text.replace(placeholder, str(value))
            return text
        except Exception as e:
            print(f"Error replacing variables: {e}")
            return text
    
    def validate_template(self) -> Dict[str, Any]:
        """Validate the template content and return validation results."""
        validation_result = {
            'is_valid': True,
            'errors': [],
            'warnings': [],
            'tag_count': 0,
            'variable_count': 0
        }
        
        try:
            # Check for basic HTML structure
            soup = BeautifulSoup(self.template_content, 'html.parser')
            validation_result['tag_count'] = len(soup.find_all())
            
            # Count variables
            import re
            variable_pattern = r'<(\w+)/>'
            variables = re.findall(variable_pattern, self.template_content)
            validation_result['variable_count'] = len(set(variables))
            
            # Check for missing variables in context
            missing_vars = []
            for var in set(variables):
                if var not in self.context:
                    missing_vars.append(var)
                    validation_result['warnings'].append(f"Variable '{var}' not found in context")
            
            if missing_vars:
                validation_result['warnings'].append(f"Missing variables: {', '.join(missing_vars)}")
            
            # Check for syntax errors
            try:
                soup = BeautifulSoup(self.template_content, 'html.parser')
            except Exception as e:
                validation_result['is_valid'] = False
                validation_result['errors'].append(f"HTML parsing error: {e}")
                
        except Exception as e:
            validation_result['is_valid'] = False
            validation_result['errors'].append(f"Validation error: {e}")
        
        return validation_result
    
    def export_template_to_html(self, output_path: str = None) -> str:
        """Export the processed template to an HTML file for preview."""
        if not output_path:
            output_path = self.output_path.replace('.docx', '.html')
        
        try:
            # Create a clean HTML version
            html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Template Preview</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; font-weight: bold; }}
        .error {{ color: red; font-style: italic; }}
        h1, h2, h3, h4, h5, h6 {{ color: #333; }}
    </style>
</head>
<body>
{self.template_content}
</body>
</html>
            """
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return output_path
            
        except Exception as e:
            print(f"Error exporting HTML: {e}")
            return None

    def _handle_base64_image(self, tag: Tag):
        """Handle base64 encoded images in the template."""
        import base64
        from io import BytesIO
        image_data = tag.get_text().strip()
        if not image_data.startswith('data:image/'):
            raise ValueError("Invalid base64 image data.")
        image_data = image_data.split(',')[1]
        # Decode the base64 image data
        if image_data:
            try:
                image_bytes = base64.b64decode(image_data)
                image_stream = BytesIO(image_bytes)
                self.document.add_picture(image_stream, width=Inches(4))
                # Clean up memory
                del image_bytes
                image_stream.close()
                gc.collect()
            except Exception as e:
                self._add_error_paragraph(f"[Error loading image {tag.name}: {e}]")

    def _handle_list(self, tag: Tag):
        """Handle <ul> or <ol> as bullet or numbered list."""
        list_type = 'bullet' if tag.name == 'ul' else 'number'

        for li in tag.find_all('li', recursive=False):
            paragraph = self.document.add_paragraph(style='List Bullet' if list_type == 'bullet' else 'List Number')
            self._handle_inline_formatting(li, paragraph)

            # Xử lý nested list (ul/ol trong li)
            for nested in li.find_all(['ul', 'ol'], recursive=False):
                self._handle_list(nested)

    def cleanup_memory(self):
        """Clean up memory and resources."""
        gc.collect()
        
    def get_memory_usage(self) -> Dict[str, Any]:
        """Get current memory usage information."""
        import psutil
        import os
        
        process = psutil.Process(os.getpid())
        memory_info = process.memory_info()
        
        return {
            'rss': memory_info.rss / 1024 / 1024,  # MB
            'vms': memory_info.vms / 1024 / 1024,  # MB
            'percent': process.memory_percent()
        }
    
    def optimize_for_large_datasets(self, chunk_size: int = 500):
        """Optimize settings for processing large datasets."""
        self.chunk_size = chunk_size
        self.max_workers = min(4, os.cpu_count() or 1)
        
    def set_table_limits(self, max_rows: int = 10000, max_columns: int = 100):
        """Set limits for table processing to prevent memory issues."""
        self.max_table_rows = max_rows
        self.max_table_columns = max_columns

    def get_performance_stats(self) -> Dict[str, Any]:
        """Get performance statistics for the current export operation."""
        import time
        if hasattr(self, '_start_time'):
            elapsed_time = time.time() - self._start_time
            memory_usage = self.get_memory_usage()
            
            return {
                'elapsed_time': elapsed_time,
                'memory_usage': memory_usage,
                'template_size': len(self.template_content),
                'context_variables': len(self.context)
            }
        return {}
    
    def start_performance_monitoring(self):
        """Start performance monitoring for the export operation."""
        import time
        self._start_time = time.time()
        self._initial_memory = self.get_memory_usage()
    
    def stop_performance_monitoring(self) -> Dict[str, Any]:
        """Stop performance monitoring and return final stats."""
        if hasattr(self, '_start_time'):
            final_stats = self.get_performance_stats()
            final_stats['peak_memory'] = self._initial_memory
            delattr(self, '_start_time')
            delattr(self, '_initial_memory')
            return final_stats
        return {}
